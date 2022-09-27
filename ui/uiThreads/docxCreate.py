from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def docxCreate(data, path, type, date, username):
    ###创建新的docx文档
    document = Document()
    ###更改文档的页面大小（下面设置的是竖版A4大小页面）
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)

    ###分别设定一级标题，二级标题以及三级标题和正文的格式
    ###设置字号大小单位是磅
    document.styles['Heading 1'].font.size = Pt(16)
    ###设置文字颜色
    document.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    ###设置字体
    document.styles['Heading 1'].font.name = '微软雅黑'

    document.styles['Heading 2'].font.size = Pt(16)
    document.styles['Heading 2'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Heading 2'].font.name = '微软雅黑'

    document.styles['Heading 3'].font.size = Pt(14)
    document.styles['Heading 3'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Heading 3'].font.name = '微软雅黑'

    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    document.styles['Normal'].font.bold = False

    ###输入一级标题，二级标题
    H1 = document.add_paragraph(style='Heading 1')
    H1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    H1.paragraph_format.line_spacing = 1.5  ###设置行间距是1.5倍行间距
    run = H1.add_run('安全检测报告单')  ###添加文字
    ###中文字体显示，需要按照下面的方法设定
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 添加段落
    paragraph = document.add_paragraph(style='Normal')
    run = paragraph.add_run('检测日期：%s \t检测负责人：%s' % (date,username))
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    paragraph.paragraph_format.line_spacing = 1.5

    # 增加表格
    row = len(data)
    col = 5
    table = document.add_table(rows=1, cols=col, style='Table Grid')  ###style可以选择不同的，具体可以百度
    hdr_cells = table.rows[0].cells
    for i in range(col):
        hdr_cells[i].text = (["编号", "物品单号", "检测违禁品", "是否违禁品", "检测时间", "检测信息"])[i]
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中

    # 再增加表格内的元素
    for i in range(row):
        tr = table.rows[i]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "300")
        trPr.append(trHeight)
        rowdata = data[i]
        row_cells = table.add_row().cells
        for j in range(col):
            if j == 1:
                temp = rowdata[j].split("/")
                row_cells[j].text = str(temp[-1])
            elif j == 3:
                if rowdata[j] == 0:
                    row_cells[j].text = "否"
                else:
                    row_cells[j].text = "是"
            else:
                row_cells[j].text = str(rowdata[j])
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直对中
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
        if i == (row - 1):
            for j in range(col):
                table.cell(0, j)._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor='D8D8D8')))  ###第一行加入底色
                table.cell(-1, j)._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor='DCE6F2')))  ###最后一行加入底色

    # 设定表格里的字体和字号
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设定每一列的宽度，对应每一列的宽度之和应该为10，然后按照比例平均分配
    col_width_1 = {0: 1.0, 1: 1, 2: 1.0, 3: 1, 4: 1, 5: 1.5, 6: 1.5, 7: 2}
    for col_num in range(col):
        table.cell(0, col_num).width = Inches(col_width_1[col_num])

    document.save(path)